# scripts/04_trial_landscape_summary.py

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt



# =========================
# 1. Define paths
# =========================

BASE_DIR = Path(__file__).resolve().parents[1]

INPUT_PATH = BASE_DIR / "data" / "processed" / "clinical_trials_octreotide_cleaned.xlsx"

REPORTS_DIR = BASE_DIR / "reports"
TABLES_DIR = REPORTS_DIR / "tables"

TABLES_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_EXCEL_PATH = TABLES_DIR / "octreotide_trial_landscape_summary.xlsx"
OUTPUT_WORD_PATH = REPORTS_DIR / "octreotide_trial_landscape_interpretation.docx"


# =========================
# 2. Load cleaned data
# =========================

print("Loading cleaned clinical trial dataset...")

df = pd.read_excel(INPUT_PATH)

print(f"Dataset loaded successfully.")
print(f"Shape: {df.shape}")


# =========================
# 3. Basic safety cleaning
# =========================

print("\nPreparing data for landscape analysis...")

if "status" in df.columns:
    df["status"] = df["status"].astype(str).str.strip().str.upper()

if "phase" in df.columns:
    df["phase"] = df["phase"].astype(str).str.strip().str.upper()
    df["phase"] = df["phase"].replace({"NAN": "MISSING", "NONE": "MISSING", "": "MISSING"})

if "sponsor" in df.columns:
    df["sponsor"] = df["sponsor"].astype(str).str.strip()
    df["sponsor"] = df["sponsor"].replace({"nan": "Missing", "None": "Missing", "": "Missing"})

if "enrollment" in df.columns:
    df["enrollment"] = pd.to_numeric(df["enrollment"], errors="coerce")

if "trial_duration_months" in df.columns:
    df["trial_duration_months"] = pd.to_numeric(df["trial_duration_months"], errors="coerce")

if "start_year" in df.columns:
    df["start_year"] = pd.to_numeric(df["start_year"], errors="coerce")


# =========================
# 4. Create active/inactive classification
# =========================

print("\nClassifying trials as active or inactive...")

active_statuses = [
    "RECRUITING",
    "NOT_YET_RECRUITING",
    "ACTIVE_NOT_RECRUITING",
    "ENROLLING_BY_INVITATION"
]

if "status" in df.columns:
    df["activity_group"] = df["status"].apply(
        lambda x: "Active / ongoing" if x in active_statuses else "Inactive / closed"
    )


# =========================
# 5. Core landscape metrics
# =========================

print("\nCalculating core landscape metrics...")

total_trials = len(df)

completed_trials = (df["status"] == "COMPLETED").sum() if "status" in df.columns else None
recruiting_trials = (df["status"] == "RECRUITING").sum() if "status" in df.columns else None
terminated_trials = (df["status"] == "TERMINATED").sum() if "status" in df.columns else None
withdrawn_trials = (df["status"] == "WITHDRAWN").sum() if "status" in df.columns else None

active_trials = (df["activity_group"] == "Active / ongoing").sum() if "activity_group" in df.columns else None
inactive_trials = (df["activity_group"] == "Inactive / closed").sum() if "activity_group" in df.columns else None

median_enrollment = df["enrollment"].median() if "enrollment" in df.columns else None
max_enrollment = df["enrollment"].max() if "enrollment" in df.columns else None

median_duration = df["trial_duration_months"].median() if "trial_duration_months" in df.columns else None
max_duration = df["trial_duration_months"].max() if "trial_duration_months" in df.columns else None

earliest_start_year = df["start_year"].min() if "start_year" in df.columns else None
latest_start_year = df["start_year"].max() if "start_year" in df.columns else None


summary_metrics = pd.DataFrame({
    "metric": [
        "Total trials",
        "Completed trials",
        "Recruiting trials",
        "Terminated trials",
        "Withdrawn trials",
        "Active / ongoing trials",
        "Inactive / closed trials",
        "Median enrollment",
        "Maximum enrollment",
        "Median trial duration, months",
        "Maximum trial duration, months",
        "Earliest start year",
        "Latest start year"
    ],
    "value": [
        total_trials,
        completed_trials,
        recruiting_trials,
        terminated_trials,
        withdrawn_trials,
        active_trials,
        inactive_trials,
        median_enrollment,
        max_enrollment,
        median_duration,
        max_duration,
        earliest_start_year,
        latest_start_year
    ]
})


# =========================
# 6. Detailed summary tables
# =========================

print("\nCreating detailed summary tables...")

if "status" in df.columns:
    status_summary = (
        df["status"]
        .value_counts(dropna=False)
        .rename_axis("status")
        .reset_index(name="count")
    )
else:
    status_summary = pd.DataFrame(columns=["status", "count"])


if "phase" in df.columns:
    phase_summary = (
        df["phase"]
        .value_counts(dropna=False)
        .rename_axis("phase")
        .reset_index(name="count")
    )
else:
    phase_summary = pd.DataFrame(columns=["phase", "count"])


if "activity_group" in df.columns:
    activity_summary = (
        df["activity_group"]
        .value_counts(dropna=False)
        .rename_axis("activity_group")
        .reset_index(name="count")
    )
else:
    activity_summary = pd.DataFrame(columns=["activity_group", "count"])


if "sponsor" in df.columns:
    top_sponsors = (
        df["sponsor"]
        .value_counts(dropna=False)
        .head(20)
        .rename_axis("sponsor")
        .reset_index(name="count")
    )
else:
    top_sponsors = pd.DataFrame(columns=["sponsor", "count"])


if "start_year" in df.columns:
    year_summary = (
        df["start_year"]
        .dropna()
        .astype(int)
        .value_counts()
        .sort_index()
        .rename_axis("start_year")
        .reset_index(name="count")
    )
else:
    year_summary = pd.DataFrame(columns=["start_year", "count"])

if {"phase", "status"}.issubset(df.columns):
    phase_status_table = pd.crosstab(
        df["phase"],
        df["status"],
        margins=True
    )
else:
    phase_status_table = pd.DataFrame()


if {"sponsor", "phase"}.issubset(df.columns):
    sponsor_phase_table = pd.crosstab(
        df["sponsor"],
        df["phase"]
    )
else:
    sponsor_phase_table = pd.DataFrame()

# =========================
# 7. Create trial-level analytical table
# =========================

print("\nCreating trial-level analytical table...")

selected_columns = [
    "nct_id",
    "title",
    "status",
    "activity_group",
    "phase",
    "sponsor",
    "enrollment",
    "trial_duration_months",
    "start_year",
    "conditions"
]

existing_selected_columns = [col for col in selected_columns if col in df.columns]

trial_level_table = df[existing_selected_columns].copy()

if "enrollment" in trial_level_table.columns:
    trial_level_table["enrollment_group"] = pd.cut(
        trial_level_table["enrollment"],
        bins=[-1, 50, 100, 500, 1000, float("inf")],
        labels=[
            "0-50",
            "51-100",
            "101-500",
            "501-1000",
            ">1000"
        ]
    )

if "trial_duration_months" in trial_level_table.columns:
    trial_level_table["duration_group"] = pd.cut(
        trial_level_table["trial_duration_months"],
        bins=[-1, 12, 24, 36, 60, float("inf")],
        labels=[
            "<=12 months",
            "13-24 months",
            "25-36 months",
            "37-60 months",
            ">60 months"
        ]
    )


# =========================
# 8. Write Excel report
# =========================

print("\nSaving Excel landscape report...")

with pd.ExcelWriter(OUTPUT_EXCEL_PATH, engine="openpyxl") as writer:
    summary_metrics.to_excel(writer, sheet_name="summary_metrics", index=False)
    status_summary.to_excel(writer, sheet_name="status_summary", index=False)
    phase_summary.to_excel(writer, sheet_name="phase_summary", index=False)
    activity_summary.to_excel(writer, sheet_name="activity_summary", index=False)
    top_sponsors.to_excel(writer, sheet_name="top_sponsors", index=False)
    year_summary.to_excel(writer, sheet_name="year_summary", index=False)
    phase_status_table.to_excel(writer, sheet_name="phase_by_status")
    sponsor_phase_table.to_excel(writer, sheet_name="sponsor_by_phase")
    trial_level_table.to_excel(writer, sheet_name="trial_level_table", index=False)

print(f"Excel report saved to: {OUTPUT_EXCEL_PATH}")


# =========================
# 9. Generate written interpretation
# =========================

print("\nGenerating written landscape interpretation...")

top_status = status_summary.iloc[0]["status"] if not status_summary.empty else "not available"
top_phase = phase_summary.iloc[0]["phase"] if not phase_summary.empty else "not available"
top_sponsor = top_sponsors.iloc[0]["sponsor"] if not top_sponsors.empty else "not available"

interpretation = f"""
Octreotide Clinical Trial Landscape Summary

Dataset overview
----------------
The dataset contains {total_trials} octreotide-related clinical trial records extracted from ClinicalTrials.gov.

Trial status landscape
----------------------
The most common trial status is {top_status}. The dataset includes {completed_trials} completed trials, {recruiting_trials} recruiting trials, {terminated_trials} terminated trials, and {withdrawn_trials} withdrawn trials.

Based on recruitment status, {active_trials} trials are classified as active or ongoing, while {inactive_trials} trials are classified as inactive or closed.

Clinical phase landscape
------------------------
The most common development phase is {top_phase}. This suggests that the octreotide clinical trial landscape includes substantial mid- to late-stage clinical development activity.

Sponsor landscape
-----------------
The most frequent sponsor in the dataset is {top_sponsor}. Sponsor-level analysis helps identify the major commercial, academic, or institutional contributors to the octreotide trial landscape.

Enrollment profile
------------------
The median enrollment size is {median_enrollment} participants, while the largest trial enrolled {max_enrollment} participants. This indicates that most octreotide studies are relatively focused in size, although some larger studies are also present.

Trial duration profile
----------------------
The median trial duration is approximately {median_duration} months, with the longest trial lasting approximately {max_duration} months. Trial duration may reflect differences in indication, phase, endpoint type, follow-up requirements, and recruitment complexity.

Time trend
----------
The available start year range is from {earliest_start_year} to {latest_start_year}. This allows investigation of whether octreotide trial activity has increased, decreased, or shifted over time.

Portfolio interpretation
------------------------
From a pharma data science perspective, this analysis provides a structured overview of the octreotide clinical development landscape. The dataset can be used to explore development trends, sponsor concentration, phase distribution, recruitment status, enrollment burden, and trial duration patterns.

Potential next steps
--------------------
1. Analyse trial duration by phase and sponsor.
2. Compare enrollment size across trial phases.
3. Identify indications most commonly associated with octreotide trials.
4. Build a recruitment-risk or trial-duration modelling dataset.
5. Expand the search to comparator products such as lanreotide or pasireotide.
"""


doc = Document()

# Main title
title = doc.add_heading("Octreotide Clinical Trial Landscape Summary", level=1)

# Add report sections
sections = {
    "Dataset overview": f"The dataset contains {total_trials} octreotide-related clinical trial records extracted from ClinicalTrials.gov.",

    "Trial status landscape": f"The most common trial status is {top_status}. The dataset includes {completed_trials} completed trials, {recruiting_trials} recruiting trials, {terminated_trials} terminated trials, and {withdrawn_trials} withdrawn trials. Based on recruitment status, {active_trials} trials are classified as active or ongoing, while {inactive_trials} trials are classified as inactive or closed.",

    "Clinical phase landscape": f"The most common development phase is {top_phase}. This suggests that the octreotide clinical trial landscape includes substantial mid- to late-stage clinical development activity.",

    "Sponsor landscape": f"The most frequent sponsor in the dataset is {top_sponsor}. Sponsor-level analysis helps identify the major commercial, academic, or institutional contributors to the octreotide trial landscape.",

    "Enrollment profile": f"The median enrollment size is {median_enrollment} participants, while the largest trial enrolled {max_enrollment} participants. This indicates that most octreotide studies are relatively focused in size, although some larger studies are also present.",

    "Trial duration profile": f"The median trial duration is approximately {median_duration} months, with the longest trial lasting approximately {max_duration} months. Trial duration may reflect differences in indication, phase, endpoint type, follow-up requirements, and recruitment complexity.",

    "Time trend": f"The available start year range is from {earliest_start_year} to {latest_start_year}. This allows investigation of whether octreotide trial activity has increased, decreased, or shifted over time.",

    "Portfolio interpretation": "From a pharma data science perspective, this analysis provides a structured overview of the octreotide clinical development landscape. The dataset can be used to explore development trends, sponsor concentration, phase distribution, recruitment status, enrollment burden, and trial duration patterns."
}

for heading, paragraph in sections.items():
    doc.add_heading(heading, level=2)
    doc.add_paragraph(paragraph)


# Add next steps
doc.add_heading("Potential next steps", level=2)

next_steps = [
    "Analyse trial duration by phase and sponsor.",
    "Compare enrollment size across trial phases.",
    "Identify indications most commonly associated with octreotide trials.",
    "Build a recruitment-risk or trial-duration modelling dataset.",
    "Expand the search to comparator products such as lanreotide or pasireotide."
]

for step in next_steps:
    doc.add_paragraph(step, style="List Bullet")


# Optional: set normal font size
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.size = Pt(11)


doc.save(OUTPUT_WORD_PATH)

print(f"Word interpretation report saved to: {OUTPUT_WORD_PATH}")
# =========================
# 10. Finish
# =========================

print("\nStage 4 clinical trial landscape summary completed successfully.")